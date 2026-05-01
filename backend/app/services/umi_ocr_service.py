"""Local Umi-OCR integration and editable Word export for resume images."""

from __future__ import annotations

import base64
import io
import re
import subprocess
import time
from pathlib import Path
from typing import Any, Dict, Iterable, List

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


WORD_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
LOW_CONFIDENCE_THRESHOLD = 0.82
SECTION_KEYWORDS = (
    "个人信息",
    "教育",
    "教育背景",
    "实习",
    "实习经历",
    "工作经历",
    "项目",
    "项目经历",
    "校园经历",
    "获奖",
    "奖项",
    "荣誉",
    "证书",
    "技能",
    "个人技能",
    "自我评价",
    "个人总结",
)


class UmiOcrService:
    """Call local Umi-OCR and convert OCR blocks into an editable DOCX."""

    def __init__(
        self,
        *,
        base_url: str,
        exe_path: Path,
        auto_start: bool,
        document_dir: Path,
        project_root: Path,
    ):
        self.base_url = (base_url or "http://127.0.0.1:1224").rstrip("/")
        self.exe_path = self._resolve_project_path(exe_path, project_root)
        self.auto_start = bool(auto_start)
        self.document_dir = self._resolve_project_path(document_dir, project_root)
        self.document_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _resolve_project_path(path: Path, project_root: Path) -> Path:
        resolved = Path(path)
        if not resolved.is_absolute():
            resolved = project_root / resolved
        return resolved.resolve()

    def probe(self) -> Dict[str, Any]:
        """Return Umi-OCR HTTP status without running OCR."""

        started_at = time.perf_counter()
        try:
            response = requests.get(f"{self.base_url}/api/ocr/get_options", timeout=2.5)
            response.raise_for_status()
            return {
                "configured": self.exe_path.exists(),
                "reachable": True,
                "status": "available",
                "latency_ms": int((time.perf_counter() - started_at) * 1000),
                "error": "",
            }
        except requests.RequestException as exc:
            return {
                "configured": self.exe_path.exists(),
                "reachable": False,
                "status": "not_running" if self.exe_path.exists() else "missing_runtime",
                "latency_ms": None,
                "error": str(exc),
            }

    def ensure_available(self) -> None:
        """Start Umi-OCR when configured and wait for the HTTP endpoint."""

        if self.probe()["reachable"]:
            return
        if not self.auto_start:
            raise RuntimeError("Umi-OCR HTTP service is not running.")
        if not self.exe_path.exists():
            raise RuntimeError(f"Umi-OCR.exe was not found: {self.exe_path}")

        creationflags = 0
        if hasattr(subprocess, "CREATE_NO_WINDOW"):
            creationflags = subprocess.CREATE_NO_WINDOW

        try:
            subprocess.Popen(
                [str(self.exe_path), "--hide"],
                cwd=str(self.exe_path.parent),
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                creationflags=creationflags,
            )
        except OSError as exc:
            raise RuntimeError(f"Failed to start Umi-OCR.exe: {exc}") from exc

        deadline = time.perf_counter() + 45
        while time.perf_counter() < deadline:
            if self.probe()["reachable"]:
                return
            time.sleep(1)

        raise RuntimeError("Umi-OCR started but the HTTP service did not become reachable.")

    def export_word_from_image(self, *, image_path: Path, file_name: str) -> Dict[str, Any]:
        """OCR one generated image and return a saved editable DOCX as bytes."""

        if not image_path.exists() or not image_path.is_file():
            raise FileNotFoundError("Generated resume image was not found.")

        self.ensure_available()
        payload = self._post_ocr(image_path)
        blocks = self._normalize_blocks(payload.get("data") or [])
        paragraphs = self._paragraphs_from_blocks(blocks)
        low_confidence_count = sum(1 for block in blocks if block.get("score", 1) < LOW_CONFIDENCE_THRESHOLD)

        docx_bytes = self._build_docx(paragraphs, blocks=blocks, low_confidence_count=low_confidence_count)
        safe_name = self._safe_docx_name(file_name)
        saved_path = self.document_dir / safe_name
        saved_path.write_bytes(docx_bytes.getvalue())
        docx_bytes.seek(0)

        return {
            "content": docx_bytes,
            "media_type": WORD_MEDIA_TYPE,
            "file_name": safe_name,
            "saved_name": safe_name,
            "block_count": len(blocks),
            "low_confidence_count": low_confidence_count,
            "ocr_time": payload.get("time"),
        }

    def get_generated_document_path(self, file_name: str) -> Path:
        """Return one generated OCR document path under the configured directory."""

        safe_name = Path(file_name or "").name
        if not safe_name or safe_name != file_name:
            raise FileNotFoundError("Generated Word document was not found.")
        path = (self.document_dir / safe_name).resolve()
        if self.document_dir.resolve() not in path.parents and path != self.document_dir.resolve():
            raise FileNotFoundError("Generated Word document was not found.")
        if not path.exists() or not path.is_file():
            raise FileNotFoundError("Generated Word document was not found.")
        return path

    def _post_ocr(self, image_path: Path) -> Dict[str, Any]:
        encoded = base64.b64encode(image_path.read_bytes()).decode("ascii")
        primary_options = {
            "ocr.language": "简体中文",
            "ocr.angle": False,
            "ocr.maxSideLen": 4096,
            "tbpu.parser": "multi_para",
            "data.format": "dict",
        }
        fallback_options = {
            "tbpu.parser": "multi_para",
            "data.format": "dict",
        }

        try:
            return self._request_ocr(encoded, primary_options)
        except RuntimeError as exc:
            # Some Umi-OCR builds expose a smaller option surface. Retry with the
            # minimal dict-return options before surfacing a hard failure.
            if "option" not in str(exc).lower() and "参数" not in str(exc):
                raise
            return self._request_ocr(encoded, fallback_options)

    def _request_ocr(self, encoded: str, options: Dict[str, Any]) -> Dict[str, Any]:
        try:
            response = requests.post(
                f"{self.base_url}/api/ocr",
                json={"base64": encoded, "options": options},
                headers={"Content-Type": "application/json"},
                timeout=120,
            )
            response.raise_for_status()
            payload = response.json()
        except requests.RequestException as exc:
            raise RuntimeError(f"Umi-OCR request failed: {exc}") from exc
        except ValueError as exc:
            raise RuntimeError("Umi-OCR returned invalid JSON.") from exc

        code = payload.get("code")
        if code == 100:
            return payload
        if code == 101:
            return {**payload, "data": []}

        message = payload.get("data") or payload.get("message") or "unknown OCR error"
        raise RuntimeError(f"Umi-OCR failed with code {code}: {message}")

    def _normalize_blocks(self, raw_blocks: Any) -> List[Dict[str, Any]]:
        if not isinstance(raw_blocks, list):
            return []

        blocks: List[Dict[str, Any]] = []
        for item in raw_blocks:
            if not isinstance(item, dict):
                continue
            text = re.sub(r"[ \t]+", " ", str(item.get("text") or "")).strip()
            if not text:
                continue
            box = item.get("box") or []
            xs, ys = self._box_axes(box)
            blocks.append(
                {
                    "text": text,
                    "score": self._score(item.get("score")),
                    "box": box,
                    "x": min(xs) if xs else 0,
                    "y": min(ys) if ys else 0,
                    "height": (max(ys) - min(ys)) if len(ys) >= 2 else 18,
                    "end": str(item.get("end") or ""),
                }
            )

        blocks.sort(key=lambda block: (block["y"], block["x"]))
        return blocks

    @staticmethod
    def _box_axes(box: Any) -> tuple[List[float], List[float]]:
        xs: List[float] = []
        ys: List[float] = []
        if not isinstance(box, list):
            return xs, ys
        for point in box:
            if isinstance(point, list) and len(point) >= 2:
                try:
                    xs.append(float(point[0]))
                    ys.append(float(point[1]))
                except (TypeError, ValueError):
                    continue
        return xs, ys

    @staticmethod
    def _score(value: Any) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return 1.0

    def _paragraphs_from_blocks(self, blocks: List[Dict[str, Any]]) -> List[str]:
        if not blocks:
            return []

        with_end = any(block.get("end") for block in blocks)
        if with_end:
            paragraphs: List[str] = []
            current = ""
            for block in blocks:
                ending = block.get("end") or " "
                current = f"{current}{block['text']}{ending}"
                if "\n" in ending:
                    for paragraph in current.splitlines():
                        cleaned = paragraph.strip()
                        if cleaned:
                            paragraphs.append(cleaned)
                    current = ""
            if current.strip():
                paragraphs.append(current.strip())
            return self._dedupe_blank_lines(paragraphs)

        rows = self._group_blocks_by_row(blocks)
        return self._dedupe_blank_lines([" ".join(block["text"] for block in row).strip() for row in rows])

    def _group_blocks_by_row(self, blocks: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
        rows: List[List[Dict[str, Any]]] = []
        median_height = sorted(block.get("height", 18) for block in blocks)[len(blocks) // 2]
        row_threshold = max(14, median_height * 0.65)

        for block in blocks:
            placed = False
            for row in rows:
                row_y = sum(item["y"] for item in row) / len(row)
                if abs(block["y"] - row_y) <= row_threshold:
                    row.append(block)
                    placed = True
                    break
            if not placed:
                rows.append([block])

        for row in rows:
            row.sort(key=lambda item: item["x"])
        rows.sort(key=lambda row: min(item["y"] for item in row))
        return rows

    @staticmethod
    def _dedupe_blank_lines(lines: Iterable[str]) -> List[str]:
        cleaned: List[str] = []
        for line in lines:
            stripped = re.sub(r"\s+", " ", line).strip()
            if not stripped:
                continue
            if cleaned and cleaned[-1] == stripped:
                continue
            cleaned.append(stripped)
        return cleaned

    def _build_docx(
        self,
        paragraphs: List[str],
        *,
        blocks: List[Dict[str, Any]],
        low_confidence_count: int,
    ) -> io.BytesIO:
        document = Document()
        section = document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(16)
        section.bottom_margin = Mm(16)
        section.left_margin = Mm(17)
        section.right_margin = Mm(17)

        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        if not paragraphs:
            paragraph = document.add_paragraph()
            paragraph.add_run("Umi-OCR 未在这张图片中识别到可导出的文字。").bold = True
        else:
            first = True
            for text in paragraphs:
                paragraph = document.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(4)
                paragraph_format.line_spacing = 1.12
                run = paragraph.add_run(text)
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

                if first and len(text) <= 24:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.bold = True
                    run.font.size = Pt(16)
                    first = False
                elif self._looks_like_section_title(text):
                    self._style_section_title(paragraph, run)
                else:
                    run.font.size = Pt(10.5)
                    first = False

        document.add_paragraph()
        note = document.add_paragraph()
        note_run = note.add_run(
            f"OCR 提示：共识别 {len(blocks)} 个文本块，低置信度 {low_confidence_count} 个。"
            "请重点检查姓名、电话、邮箱、时间、项目名和数字。"
        )
        note_run.font.size = Pt(8.5)
        note_run.font.color.rgb = RGBColor(100, 116, 139)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _looks_like_section_title(text: str) -> bool:
        compact = re.sub(r"[\s:：|｜/\\-]+", "", text)
        if len(compact) > 16:
            return False
        return any(keyword in compact for keyword in SECTION_KEYWORDS)

    @staticmethod
    def _style_section_title(paragraph, run) -> None:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 64, 175)

        p_pr = paragraph._p.get_or_add_pPr()
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D8E2F3")
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        borders.append(bottom)

    @staticmethod
    def _safe_docx_name(file_name: str) -> str:
        original = Path(file_name or "resume-image-ocr").stem
        safe_stem = re.sub(r"[^A-Za-z0-9_\-\u4e00-\u9fff]+", "_", original)
        safe_stem = re.sub(r"_+", "_", safe_stem).strip("._- ") or "resume-image-ocr"
        return f"{safe_stem}.docx"
